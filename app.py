"""
SmartNote - AI powered study assistant
Features: syllabus summarization, quiz generation, answer sheet grading, voice assistant
Uses the free Groq API for all LLM calls.
"""

import os
import json
import re
import uuid
import base64

from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import requests

import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
# llama-3.3-70b-versatile was deprecated by Groq — gpt-oss-120b is the recommended replacement.
GROQ_MODEL = os.getenv("GROQ_MODEL", "openai/gpt-oss-120b")
# Vision-capable model, used for handwriting OCR and any image reasoning.
GROQ_VISION_MODEL = os.getenv("GROQ_VISION_MODEL", "qwen/qwen3.6-27b")
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "uploads")
ALLOWED_EXTENSIONS = {"pdf", "docx", "jpg", "jpeg", "png", "txt"}
MAX_CONTENT_LENGTH = 20 * 1024 * 1024  # 20 MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath: str) -> str:
    text_parts = []
    with fitz.open(filepath) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts).strip()


def extract_text_from_image(filepath: str, handwritten: bool = False) -> str:
    """
    Typed/printed text goes through Tesseract (fast, free, local).
    Handwritten notes go through a Groq vision model instead, since Tesseract's
    handwriting accuracy is too low to be usable. We also fall back to the
    vision model automatically if Tesseract comes back empty/near-empty.
    """
    if not handwritten:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image).strip()
        if len(text) >= 15:
            return text
        # Tesseract found little or nothing — likely handwriting or a low quality
        # scan. Try the vision model before giving up.

    return extract_text_from_image_vision(filepath)


def extract_text_from_image_vision(filepath: str) -> str:
    with open(filepath, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")

    ext = filepath.rsplit(".", 1)[-1].lower()
    mime = "image/png" if ext == "png" else "image/jpeg"

    messages = [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Transcribe every word of text in this image exactly as written, "
                        "including handwriting. Preserve line breaks and structure "
                        "(headings, bullet points, numbered lists) where visible. "
                        "Output only the transcribed text — no commentary, no preamble, "
                        "no markdown code fences. If part of the writing is illegible, "
                        "mark that spot with [illegible] rather than guessing."
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                },
            ],
        }
    ]
    text = call_groq(messages, temperature=0.1, max_tokens=4000, model=GROQ_VISION_MODEL)
    return strip_code_fence(text).strip()


def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def extract_text(filepath: str, filename: str, handwritten: bool = False) -> str:
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    if ext == "docx":
        return extract_text_from_docx(filepath)
    if ext in ("jpg", "jpeg", "png"):
        return extract_text_from_image(filepath, handwritten=handwritten)
    if ext == "txt":
        return extract_text_from_txt(filepath)
    raise ValueError(f"Unsupported file type: {ext}")


def call_groq(messages, temperature=0.4, max_tokens=3000, json_mode=False, model=None):
    if not GROQ_API_KEY:
        raise RuntimeError(
            "GROQ_API_KEY is not set. Add it to your .env file (see .env.example)."
        )

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model or GROQ_MODEL,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    if json_mode:
        payload["response_format"] = {"type": "json_object"}

    resp = requests.post(GROQ_URL, headers=headers, json=payload, timeout=60)
    resp.raise_for_status()
    data = resp.json()
    return data["choices"][0]["message"]["content"]


def strip_code_fence(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def chunk_text(text: str, max_chars: int = 12000):
    """Groq context is generous, but keep requests reasonably sized."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    for i in range(0, len(text), max_chars):
        chunks.append(text[i:i + max_chars])
    return chunks


# --------------------------------------------------------------------------
# Page routes
# --------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


# --------------------------------------------------------------------------
# API: file upload / text extraction
# --------------------------------------------------------------------------

@app.route("/api/extract", methods=["POST"])
def api_extract():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type"}), 400

    handwritten = request.form.get("handwritten", "false").lower() == "true"

    filename = secure_filename(file.filename)
    unique_name = f"{uuid.uuid4().hex}_{filename}"
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
    file.save(filepath)

    try:
        text = extract_text(filepath, filename, handwritten=handwritten)
    except Exception as e:
        return jsonify({"error": f"Failed to extract text: {e}"}), 500
    finally:
        # Clean up the uploaded file, we only need the extracted text
        if os.path.exists(filepath):
            os.remove(filepath)

    if not text:
        return jsonify({"error": "No readable text found in the file"}), 422

    return jsonify({"text": text, "filename": filename, "char_count": len(text)})


# --------------------------------------------------------------------------
# API: summarization / short notes
# --------------------------------------------------------------------------

@app.route("/api/summarize", methods=["POST"])
def api_summarize():
    data = request.get_json(force=True)
    text = (data or {}).get("text", "").strip()
    if not text:
        return jsonify({"error": "No text provided"}), 400

    chunks = chunk_text(text)
    partial_notes = []

    for chunk in chunks:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are SmartNote, an assistant that converts syllabus or study "
                    "material into clear, short, exam-friendly notes. Use markdown: "
                    "headings for topics, bullet points for key facts, and bold for "
                    "important terms. Be concise but keep all important concepts."
                ),
            },
            {
                "role": "user",
                "content": f"Summarize the following study material into short notes:\n\n{chunk}",
            },
        ]
        partial_notes.append(call_groq(messages, temperature=0.3))

    notes = "\n\n".join(partial_notes)

    # If we had to chunk, do a final pass to merge/clean up the combined notes
    if len(chunks) > 1:
        merge_messages = [
            {
                "role": "system",
                "content": (
                    "You merge multiple partial note sets into one clean, "
                    "well-organized set of short notes in markdown, removing "
                    "duplicate points."
                ),
            },
            {"role": "user", "content": notes},
        ]
        notes = call_groq(merge_messages, temperature=0.3)

    return jsonify({"notes": notes})


# --------------------------------------------------------------------------
# API: quiz generation
# --------------------------------------------------------------------------

@app.route("/api/quiz", methods=["POST"])
def api_quiz():
    data = request.get_json(force=True)
    text = (data or {}).get("text", "").strip()
    num_questions = int((data or {}).get("num_questions", 5))
    num_questions = max(1, min(num_questions, 20))

    if not text:
        return jsonify({"error": "No text provided"}), 400

    chunks = chunk_text(text, max_chars=12000)
    source_text = chunks[0]  # quiz generation works best from a single coherent pass

    messages = [
        {
            "role": "system",
            "content": (
                "You generate multiple-choice quiz questions from study material. "
                "Respond ONLY with valid JSON, no prose, no markdown code fences, "
                "in this exact shape:\n"
                '{"questions": [{"question": "...", "options": ["A", "B", "C", "D"], '
                '"correct_index": 0, "explanation": "..."}]}'
            ),
        },
        {
            "role": "user",
            "content": (
                f"Create {num_questions} multiple-choice questions from this material. "
                f"Vary difficulty and cover different topics:\n\n{source_text}"
            ),
        },
    ]

    raw = call_groq(messages, temperature=0.5, json_mode=True)

    try:
        parsed = json.loads(strip_code_fence(raw))
    except json.JSONDecodeError:
        return jsonify({"error": "Failed to parse quiz from AI response", "raw": raw}), 500

    questions = parsed.get("questions", [])
    if not questions:
        return jsonify({"error": "AI returned no questions", "raw": raw}), 500

    return jsonify({"questions": questions})


# --------------------------------------------------------------------------
# API: answer sheet grading
# --------------------------------------------------------------------------

@app.route("/api/grade", methods=["POST"])
def api_grade():
    data = request.get_json(force=True)
    questions_text = (data or {}).get("questions_text", "").strip()
    answer_text = (data or {}).get("answer_text", "").strip()
    max_marks_per_q = (data or {}).get("max_marks_per_question", 5)

    if not questions_text or not answer_text:
        return jsonify({"error": "Both questions_text and answer_text are required"}), 400

    messages = [
        {
            "role": "system",
            "content": (
                "You are an exam evaluator. You will be given a question paper (or "
                "answer key) and a student's answer sheet. Match each answer to its "
                "question, evaluate correctness and completeness, and assign marks "
                f"out of {max_marks_per_q} per question. Be fair but rigorous. "
                "Respond ONLY with valid JSON, no prose, no markdown code fences, "
                "in this exact shape:\n"
                '{"results": [{"question_number": 1, "question": "...", '
                '"student_answer": "...", "marks_awarded": 0, "max_marks": '
                f'{max_marks_per_q}, "feedback": "..."}}], '
                '"total_awarded": 0, "total_possible": 0, "overall_feedback": "..."}'
            ),
        },
        {
            "role": "user",
            "content": (
                f"QUESTION PAPER / ANSWER KEY:\n{questions_text}\n\n"
                f"STUDENT ANSWER SHEET:\n{answer_text}\n\n"
                "Grade the student's answer sheet against the question paper."
            ),
        },
    ]

    raw = call_groq(messages, temperature=0.2, max_tokens=4000, json_mode=True)

    try:
        parsed = json.loads(strip_code_fence(raw))
    except json.JSONDecodeError:
        return jsonify({"error": "Failed to parse grading result", "raw": raw}), 500

    return jsonify(parsed)


# --------------------------------------------------------------------------
# API: chatbot
# --------------------------------------------------------------------------

MAX_CHAT_HISTORY = 20  # messages kept per request, to bound tokens/cost


@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json(force=True) or {}
    history = data.get("messages", [])
    context_text = (data.get("context") or "").strip()

    if not isinstance(history, list) or not history:
        return jsonify({"error": "No messages provided"}), 400

    history = history[-MAX_CHAT_HISTORY:]

    system_content = (
        "You are the SmartNote study assistant chatbot. Help the student understand "
        "their material, answer questions, explain concepts, and give study tips. "
        "Keep answers focused and easy to follow. Use markdown (headings, bold, "
        "bullet points) when it helps readability."
    )
    if context_text:
        system_content += (
            "\n\nThe student has uploaded the following study material. Prefer "
            "grounding your answers in it when relevant, but you can also answer "
            "general questions:\n\n" + context_text[:12000]
        )

    messages = [{"role": "system", "content": system_content}]
    for m in history:
        role = m.get("role")
        content = (m.get("content") or "").strip()
        if role in ("user", "assistant") and content:
            messages.append({"role": role, "content": content})

    if len(messages) < 2:
        return jsonify({"error": "No valid messages provided"}), 400

    try:
        reply = call_groq(messages, temperature=0.5, max_tokens=1500)
    except Exception as e:
        return jsonify({"error": f"Chat failed: {e}"}), 500

    return jsonify({"reply": reply})


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
